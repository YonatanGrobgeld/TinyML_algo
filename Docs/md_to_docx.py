#!/usr/bin/env python3
"""Convert Project_Report.md into a styled Word .docx.

Figure/graph ASCII blocks are replaced by a bold TBD placeholder plus a
description of the diagram/graph that needs to be inserted.
"""
import re, os, glob
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

SRC = "/home/ronweinstein/workspace/university/Final_Project/TinyFormer_algo/Docs/Project_Report.md"
OUT = "/home/ronweinstein/workspace/university/Final_Project/TinyFormer_algo/Docs/Project_Report.docx"
FIGDIR = "/home/ronweinstein/workspace/university/Final_Project/TinyFormer_algo/Docs/figures"

WAVEDIR = "/home/ronweinstein/workspace/university/Final_Project/TinyFormer_algo/Docs/waveforms"
# Document-order figure map (titles are carried by the Word captions, not the images).
PNG = {
    1:  os.path.join(FIGDIR, "Figure1_System_Block_Diagram.png"),
    2:  os.path.join(FIGDIR, "Figure2_LiteX_SoC_Architecture.jpeg"),
    3:  os.path.join(FIGDIR, "Figure3_DOT8_Pipeline.png"),
    4:  os.path.join(FIGDIR, "Figure4_EXP_LUT_Interface.png"),
    5:  os.path.join(FIGDIR, "Figure5_GEMV_Dataflow.jpeg"),
    6:  os.path.join(FIGDIR, "Figure6_Encoder_Architecture.jpeg"),
    7:  os.path.join(WAVEDIR, "GEMV_waveform/GEMV1.jpeg"),
    8:  os.path.join(WAVEDIR, "GEMV_waveform/GEMV2.jpeg"),
    9:  os.path.join(WAVEDIR, "GEMV_waveform/GEMV3.jpeg"),
    10: os.path.join(WAVEDIR, "GEMV_waveform/GEMV4.jpeg"),
    11: os.path.join(WAVEDIR, "LUT_waveform/LUT1.jpeg"),
    12: os.path.join(WAVEDIR, "LUT_waveform/LUT2.jpeg"),
    13: os.path.join(FIGDIR, "Figure7_Baseline_Cycle_Breakdown.png"),
    14: os.path.join(FIGDIR, "Figure8_Latency_Comparison.png"),
}
# Every figure now shows its markdown caption above the image (titles were removed
# from the rendered images), so none are treated as "baked".
TITLE_BAKED = set()

# What each figure should become in the real document.
FIG = {
    1: ("Diagram", "Top-level system block diagram of the Nexys4DDR FPGA: the VexRiscv "
        "RV32IM CPU (with the DOT8 plugin) connected through the LiteX SoC bus to the "
        "EXP-LUT and GEMV MMIO peripherals, DDR2 SDRAM, and the UART. Indicate the "
        "compile-time firmware-mode selection and the serial output path for checksums "
        "and results."),
    2: ("Diagram", "Architecture diagram of the LiteX-generated SoC: VexRiscv core, the "
        "Wishbone/CSR peripheral bus, UART and Timer0, the EXP-LUT and GEMV CSR "
        "peripherals, and the DDR2 memory controller."),
    3: ("Diagram", "Datapath diagram of the VexRiscv pipeline (Fetch / Decode / Execute / "
        "Writeback) highlighting the Dot8 plugin in the Execute stage: four parallel "
        "signed int8 multipliers (4 DSP blocks) feeding an adder tree that produces the "
        "single-cycle int32 result."),
    4: ("Diagram", "Interface diagram of the EXP-LUT peripheral: the CPU writes a 4-bit "
        "index to the index CSR; a 16-entry Q10 ROM returns the corresponding exp value "
        "combinationally on the value CSR."),
    5: ("Diagram", "Dataflow diagram of the GEMV peripheral: the CPU packs four int8 lanes "
        "per 32-bit CSR write into the X and W word memories; the FSM streams one packed "
        "word from each per clock cycle through the 4-lane MAC and accumulator; results "
        "are read back via Y_OUT with Y_NEXT advancing the read pointer."),
    6: ("Diagram", "Flow diagram of the TinyFormer encoder pipeline: input [16x32] -> "
        "Q/K/V linear projections -> streaming scaled dot-product attention with softmax "
        "(EXP-LUT) -> output projection + residual -> two-layer FFN with ReLU + residual "
        "-> output [16x32]."),
    7: ("Graph", "Pie or horizontal-bar chart of the baseline 75.9 M-cycle breakdown: "
        "softmax exp() 71%, matrix-vector multiplications 21%, UART 5%, attention dot "
        "products 1%, misc 2% (data in Table 7)."),
    8: ("Graph", "Bar chart comparing per-inference latency across modes: Baseline "
        "759.00 ms (1.00x), accel_all v1 190.67 ms (3.98x), accel_all v2 157.55 ms "
        "(4.82x) (data in Table 8)."),
    9: ("Graph", "Bar chart comparing accel_all v1 (19.1 M cycles, 3.98x) against v2 "
        "(15.8 M cycles, 4.82x), annotating the ~3.3 M-cycle saving from the 4-lane "
        "packed GEMV upgrade."),
}

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
# Compact spacing so the report stays near the ~25-page target.
_npf = normal.paragraph_format
_npf.space_before = Pt(0)
_npf.space_after = Pt(4)
_npf.line_spacing = 1.05

# --- Title page (mirrors Template4ProjectReport.docx cover layout) ---
COVER = {
    "title": "TinyFormer: Hardware-Accelerated Transformer Inference on FPGA",
    "project_no": "[Project Number]",
    "report_type": "Final Project Report",
    "students": [("Ron Weinstein", "[ID]"), ("Yehonatan Grobgeld", "[ID]")],
    "supervisor": "Oren Ganon",
    "institution": "Faculty of Engineering, Tel Aviv University",
    "date": "June 2026",
}


def _center(text, size=11, bold=False, before=0, after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def build_cover():
    # vertical lead so the block sits in the upper-middle of the page
    for _ in range(3):
        doc.add_paragraph()
    _center(COVER["title"], size=22, bold=True, after=18)
    _center(f"Project No. {COVER['project_no']}", size=13, after=4)
    _center(COVER["report_type"], size=15, bold=True, after=28)

    _center("Submitted by:", size=12, bold=True, after=4)
    for name, sid in COVER["students"]:
        _center(f"{name}   —   {sid}", size=12, after=2)

    _center("Supervisor:", size=12, bold=True, before=18, after=4)
    _center(COVER["supervisor"], size=12, after=2)

    _center("Project carried out at:", size=12, bold=True, before=18, after=4)
    _center(COVER["institution"], size=12, after=2)

    _center(COVER["date"], size=12, before=28)
    doc.add_page_break()


INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+?`)")


def add_runs(paragraph, text):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        else:
            r = paragraph.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_table(rows):
    # drop a fully empty header row (e.g. the cover "| | |")
    if rows and all(c.strip() == "" for c in rows[0]):
        rows = rows[1:]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncols):
            txt = row[ci] if ci < len(row) else ""
            cell = cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, txt)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
                shade_cell(cell, "D9E2F3")
    doc.add_paragraph()


def add_code(lines):
    p = doc.add_paragraph()
    for i, ln in enumerate(lines):
        r = p.add_run(ln)
        r.font.name = "Consolas"; r.font.size = Pt(9)
        if i < len(lines) - 1:
            r.add_break()


def add_tbd(kind, desc):
    p = doc.add_paragraph()
    r = p.add_run(f"TBD - need to add {kind}: ")
    r.bold = True; r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    add_runs(p, desc)
    doc.add_paragraph()


# Waveform captures (under WAVEDIR) need to stay large enough to read;
# only the diagrams/graphs get the tighter cap.
WAVEFORM_NUMS = {n for n, p in PNG.items() if p.startswith(WAVEDIR)}


def add_image(num):
    path = PNG.get(num)
    if not path or not os.path.exists(path):
        return False
    with Image.open(path) as im:
        w, h = im.size
    aspect = w / h
    if num in WAVEFORM_NUMS:
        maxw, maxh = 5.6, 6.0      # original size — keep waveforms readable
    else:
        maxw, maxh = 5.2, 4.2      # tighter cap for diagrams/graphs
    width = maxw if maxw / aspect <= maxh else maxh * aspect
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    doc.add_paragraph()
    return True


def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click here and choose 'Update Field' to build the Table of Contents."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    r = run._r
    for e in (f1, instr, f2, t, f3):
        r.append(e)


CAP = re.compile(r"^\*+\s*Figure\s+(\d+)\b")
SEP = re.compile(r"^\s*:?-{2,}:?\s*$")

with open(SRC, encoding="utf-8") as fh:
    lines = fh.read().split("\n")

i = 0
pending_fig = None
first_title_done = False
skip_cover_table = False
# Cover handles the cover->TOC break; only Abstract needs an explicit break.
PAGEBREAK_BEFORE = {"Abstract"}

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # blank
    if stripped == "":
        i += 1
        continue

    # horizontal rule
    if stripped == "---":
        i += 1
        continue

    # headings
    hm = re.match(r"^(#{1,4})\s+(.*)$", line)
    if hm:
        hashes, htext = hm.group(1), hm.group(2).strip()
        if len(hashes) == 1 and not first_title_done:
            build_cover()
            first_title_done = True
            skip_cover_table = True
        else:
            if htext in PAGEBREAK_BEFORE:
                doc.add_page_break()
            level = max(1, len(hashes) - 1)
            doc.add_heading(htext, level=level)
            # Table of Contents -> insert a real Word TOC field, skip manual list
            if htext == "Table of Contents":
                add_toc()
                i += 1
                while i < len(lines) and not lines[i].startswith("## "):
                    i += 1
                continue
        i += 1
        continue

    # figure caption — defer; if we have a rendered PNG, the title is baked in,
    # so we skip the duplicate text caption and just insert the image.
    cm = CAP.match(stripped)
    if cm:
        num = int(cm.group(1))
        captext = stripped.strip("*").strip()
        if num not in TITLE_BAKED:
            cp = doc.add_paragraph()
            cr = cp.add_run(captext); cr.bold = True; cr.italic = True
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pending_fig = num
        i += 1
        continue

    # fenced code block
    if stripped.startswith("```"):
        i += 1
        block = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            block.append(lines[i])
            i += 1
        i += 1  # skip closing fence
        if pending_fig is not None and pending_fig in PNG:
            add_image(pending_fig)
            pending_fig = None
        elif pending_fig is not None and pending_fig in FIG:
            kind, desc = FIG[pending_fig]
            add_tbd(kind, desc)
            pending_fig = None
        else:
            add_code(block)
        continue

    # table block
    if stripped.startswith("|"):
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            raw = lines[i].strip()
            cells = [c.strip() for c in raw.strip("|").split("|")]
            if not all(SEP.match(c) or c == "" for c in cells) or not any(SEP.match(c) for c in cells):
                rows.append(cells)
            i += 1
        # The markdown cover table is now rendered by build_cover(); skip it.
        if skip_cover_table:
            skip_cover_table = False
            continue
        add_table(rows)
        continue

    # bullet list (allow indented sub-bullets)
    bm = re.match(r"^(\s*)-\s+(.*)$", line)
    if bm:
        indent = len(bm.group(1))
        body = bm.group(2)
        style = "List Bullet" if indent < 2 else "List Bullet 2"
        p = doc.add_paragraph(style=style)
        add_runs(p, body)
        i += 1
        continue

    # numbered list -> keep literal numbering, indented
    nm = re.match(r"^\s*\d+\.\s+.*$", line)
    if nm:
        p = doc.add_paragraph(style="List Paragraph")
        add_runs(p, stripped)
        i += 1
        continue

    # plain paragraph
    p = doc.add_paragraph()
    add_runs(p, stripped)
    i += 1

# Tell Word to rebuild fields (the Table of Contents) when the document is opened.
_upd = OxmlElement("w:updateFields"); _upd.set(qn("w:val"), "true")
doc.settings.element.append(_upd)

doc.save(OUT)
print("Saved", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
